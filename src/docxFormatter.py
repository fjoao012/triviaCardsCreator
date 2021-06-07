import argparse
import math
import os

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm
from utils.utils import Utils


def parse_arguments():
    parser = argparse.ArgumentParser()
    c_help = "Folder containing the cards to print in PNG format"
    parser.add_argument("-c", "--cards", help=c_help, required=True,
                        type=lambda x: Utils.is_valid_folder(parser, x))
    o_help = "Output DOCX file with all images inside, ready to print"
    parser.add_argument("-o", "--output", help=o_help, required=True,
                        type=lambda x: Utils.is_valid_file_with_extension(
                            parser, x, ".docx"))
    return parser.parse_args()


def change_orientation(document):
    for section in document.sections:
        # page orientation
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # page margins
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)


def create_document():
    document = Document()
    change_orientation(document)
    return document


def create_table_and_return_cells(document, rows, cols):
    table = document.add_table(rows=rows, cols=cols)
    table.autofit = True
    return table


def add_all_images_to_document(document, images_location, output_file):
    cells = ""
    table = ""
    rows = 2
    cols = 3
    max = rows * cols
    all_images = [x for x in os.listdir(images_location) if x.endswith(".png")]

    for index, filename in enumerate(all_images):
        if index % max == 0:
            table = create_table_and_return_cells(document, rows, cols)
            table.columns.width = Cm(3)
        pic_path = os.path.abspath(images_location + "/" + filename)
        paragraph = table.cell(math.floor((index % max) / cols), (index % max) %
                               cols).paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(pic_path, width=Cm(7.3))
    document.save(output_file)


if __name__ == "__main__":
    args = parse_arguments()
    document = create_document()
    document.save(args.output)
    add_all_images_to_document(document, args.cards, args.output)
